"""
Contract Reviewer
Upload a contract and get detailed feedback based on Voyage Advisory's contract standards
"""

import streamlit as st
import pandas as pd
import requests
import sys
import os
import json
import tempfile
import subprocess
from io import BytesIO
from datetime import datetime

# Authentication check
if 'authenticated' not in st.session_state or not st.session_state.authenticated:
    st.error("🔐 Please log in through the Home page")
    st.stop()

st.set_page_config(page_title="Contract Reviewer", page_icon="📝", layout="wide")

st.title("📝 Contract Reviewer")
st.markdown("Upload a contract for review against Voyage Advisory's contract standards")

# Config from secrets
try:
    CLAUDE_API_KEY = st.secrets.get("ANTHROPIC_API_KEY") or st.secrets.get("CLAUDE_API_KEY")
    CONTRACT_STANDARDS_DOC_ID = st.secrets.get("CONTRACT_STANDARDS_DOC_ID", "1RbPIYVgYH1HZ-FQTHYbQWycHshe-K_L5OZkat45VQnQ")
except:
    CLAUDE_API_KEY = None
    CONTRACT_STANDARDS_DOC_ID = "1RbPIYVgYH1HZ-FQTHYbQWycHshe-K_L5OZkat45VQnQ"

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def get_google_credentials():
    """Get Google credentials from Streamlit secrets"""
    try:
        from google.oauth2.service_account import Credentials
        
        # Build credentials dict from secrets
        creds_dict = {
            "type": "service_account",
            "project_id": st.secrets["gcp_service_account"]["project_id"],
            "private_key_id": st.secrets["gcp_service_account"]["private_key_id"],
            "private_key": st.secrets["gcp_service_account"]["private_key"],
            "client_email": st.secrets["gcp_service_account"]["client_email"],
            "client_id": st.secrets["gcp_service_account"]["client_id"],
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
            "client_x509_cert_url": st.secrets["gcp_service_account"].get("client_x509_cert_url", "")
        }
        
        scopes = [
            'https://www.googleapis.com/auth/documents.readonly',
            'https://www.googleapis.com/auth/drive.readonly'
        ]
        
        credentials = Credentials.from_service_account_info(creds_dict, scopes=scopes)
        return credentials
    except Exception as e:
        st.warning(f"Could not load Google credentials: {e}")
        return None

def fetch_google_doc_content(doc_id):
    """Fetch content from a Google Doc using authenticated API or public export"""
    
    # Try authenticated access first
    credentials = get_google_credentials()
    
    if credentials:
        try:
            from googleapiclient.discovery import build
            
            # Use Drive API to export as plain text (more reliable than Docs API for text extraction)
            drive_service = build('drive', 'v3', credentials=credentials)
            
            # Export as plain text
            request = drive_service.files().export_media(
                fileId=doc_id,
                mimeType='text/plain'
            )
            
            content = request.execute()
            
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            return content
            
        except Exception as e:
            st.warning(f"Authenticated access failed, trying public export: {e}")
    
    # Fallback to public export (if doc is shared publicly)
    try:
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
        response = requests.get(export_url, timeout=30)
        
        if response.status_code == 200:
            return response.text
        else:
            st.error(f"Failed to fetch Google Doc: HTTP {response.status_code}")
            return None
    except Exception as e:
        st.error(f"Error fetching Google Doc: {e}")
        return None

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF using PyPDF2"""
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            st.error("Could not extract text from PDF (may be scanned/image-based)")
            return None
            
    except ImportError:
        st.error("PyPDF2 not installed. Please add 'PyPDF2' to requirements.txt")
        return None
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
        return None

def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX using python-docx"""
    try:
        from docx import Document
        
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        if text_parts:
            return "\n\n".join(text_parts)
        else:
            st.error("Could not extract text from DOCX")
            return None
            
    except ImportError:
        st.error("python-docx not installed. Please add 'python-docx' to requirements.txt")
        return None
    except Exception as e:
        st.error(f"Error extracting DOCX text: {e}")
        return None

def extract_text_from_doc(doc_file):
    """Extract text from uploaded DOC (legacy format)"""
    try:
        # For .doc files, we need to try different approaches
        # First, try reading as if it might be a docx renamed
        try:
            from docx import Document
            doc = Document(doc_file)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            if text_parts:
                return "\n\n".join(text_parts)
        except:
            pass
        
        # Reset file position
        doc_file.seek(0)
        
        # Try reading raw text (works for some older formats)
        try:
            content = doc_file.read()
            # Try to decode as text, extracting readable portions
            if isinstance(content, bytes):
                # Extract ASCII text portions
                import re
                text = content.decode('latin-1', errors='ignore')
                # Remove control characters but keep newlines
                text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', text)
                # Clean up multiple spaces
                text = re.sub(r' +', ' ', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                if len(text.strip()) > 100:
                    return text.strip()
        except:
            pass
        
        st.error("Could not extract text from .doc file. Please convert to .docx or PDF format.")
        return None
        
    except Exception as e:
        st.error(f"Error extracting DOC text: {e}")
        return None

def call_claude_api(contract_text, standards_text):
    """Call Claude API to review contract"""
    if not CLAUDE_API_KEY:
        st.error("❌ Claude API key not configured")
        return None
    
    prompt = f"""You are a legal contract reviewer for Voyage Advisory LLC. Your task is to review the contract provided below against Voyage's contract standards.

## Voyage Contract Standards
{standards_text}

## Contract to Review
{contract_text}

## Instructions
Please provide a comprehensive review of the contract following this format:

### GENERAL COMMENTS
Provide 2-4 paragraphs summarizing the overall contract, key concerns, and general observations. Do NOT use bullet points in this section.

### DETAILED FINDINGS

For each issue found, provide a bullet point in this format:
• **Section [X.X]: [Issue Title]** - [Description of the issue and why it's a concern based on Voyage's standards]. 
  **Proposed Language:** "[Exact replacement language to use]"

For missing provisions that should be added:
• **Proposed New Section [X.X]: [Title]** - [Explanation of what's missing and why it should be added].
  **Proposed Language:** "[Exact language to insert]"

Group your findings by category:
1. Limitation of Liability
2. Work Product and Intellectual Property  
3. Payment Terms
4. Indemnification
5. Confidentiality
6. Termination
7. Governing Law and Venue
8. Entity Names and Signature Blocks
9. Other Concerns

### SUMMARY
A brief paragraph summarizing the most critical items that must be addressed before signing.

Remember:
- Reference specific section numbers from the contract
- Propose specific replacement language using Voyage's preferred standards
- Flag any unusual, non-standard, or one-sided provisions
- Check entity names in preamble and signature blocks
- Verify master agreement references if this is a SOW"""

    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": CLAUDE_API_KEY,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json"
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 8000,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=120
        )
        
        if response.status_code == 200:
            data = response.json()
            return data['content'][0]['text']
        else:
            st.error(f"Claude API error: {response.status_code} - {response.text}")
            return None
    except Exception as e:
        st.error(f"Error calling Claude API: {e}")
        return None

def create_review_docx(review_text, contract_name):
    """Create a DOCX file from the review text"""
    try:
        # Create JavaScript file for docx generation
        js_code = '''
const { Document, Packer, Paragraph, TextRun, HeadingLevel, AlignmentType, LevelFormat } = require('docx');
const fs = require('fs');

const reviewText = fs.readFileSync('/tmp/review_content.txt', 'utf8');
const contractName = fs.readFileSync('/tmp/contract_name.txt', 'utf8').trim();

// Parse the review text into sections
const lines = reviewText.split('\\n');
const children = [];

// Title
children.push(new Paragraph({
    heading: HeadingLevel.TITLE,
    children: [new TextRun({ text: "Contract Review Report", bold: true })]
}));

children.push(new Paragraph({
    children: [new TextRun({ text: `Contract: ${contractName}`, italics: true })]
}));

children.push(new Paragraph({
    children: [new TextRun({ text: `Review Date: ${new Date().toLocaleDateString()}`, italics: true })]
}));

children.push(new Paragraph({ children: [] })); // Blank line

// Process content
let inBulletSection = false;
let currentSection = '';

for (const line of lines) {
    const trimmedLine = line.trim();
    
    if (!trimmedLine) {
        children.push(new Paragraph({ children: [] }));
        continue;
    }
    
    // Main headings (###)
    if (trimmedLine.startsWith('### ')) {
        const headingText = trimmedLine.replace('### ', '').replace(/\\*\\*/g, '');
        children.push(new Paragraph({
            heading: HeadingLevel.HEADING_1,
            spacing: { before: 400, after: 200 },
            children: [new TextRun({ text: headingText, bold: true })]
        }));
        currentSection = headingText;
        continue;
    }
    
    // Sub-headings (numbered like 1. 2. etc)
    if (/^\\d+\\.\\s/.test(trimmedLine)) {
        const subHeadingText = trimmedLine.replace(/^\\d+\\.\\s/, '');
        children.push(new Paragraph({
            heading: HeadingLevel.HEADING_2,
            spacing: { before: 300, after: 150 },
            children: [new TextRun({ text: trimmedLine, bold: true })]
        }));
        continue;
    }
    
    // Bullet points
    if (trimmedLine.startsWith('• ') || trimmedLine.startsWith('- ')) {
        const bulletText = trimmedLine.replace(/^[•-]\\s/, '');
        
        // Parse bold sections
        const parts = [];
        let remaining = bulletText;
        const boldRegex = /\\*\\*([^*]+)\\*\\*/g;
        let lastIndex = 0;
        let match;
        
        while ((match = boldRegex.exec(bulletText)) !== null) {
            if (match.index > lastIndex) {
                parts.push(new TextRun({ text: bulletText.substring(lastIndex, match.index) }));
            }
            parts.push(new TextRun({ text: match[1], bold: true }));
            lastIndex = boldRegex.lastIndex;
        }
        if (lastIndex < bulletText.length) {
            parts.push(new TextRun({ text: bulletText.substring(lastIndex) }));
        }
        
        if (parts.length === 0) {
            parts.push(new TextRun({ text: bulletText }));
        }
        
        children.push(new Paragraph({
            bullet: { level: 0 },
            spacing: { before: 100, after: 100 },
            children: parts
        }));
        continue;
    }
    
    // Proposed Language (indented)
    if (trimmedLine.startsWith('**Proposed Language:**') || trimmedLine.startsWith('Proposed Language:')) {
        const langText = trimmedLine.replace(/\\*\\*Proposed Language:\\*\\*|Proposed Language:/, '').trim();
        children.push(new Paragraph({
            indent: { left: 720 },
            spacing: { before: 50, after: 150 },
            children: [
                new TextRun({ text: "Proposed Language: ", bold: true }),
                new TextRun({ text: langText, italics: true })
            ]
        }));
        continue;
    }
    
    // Regular paragraph - parse bold
    const parts = [];
    let remaining = trimmedLine;
    const boldRegex = /\\*\\*([^*]+)\\*\\*/g;
    let lastIndex = 0;
    let match;
    
    while ((match = boldRegex.exec(trimmedLine)) !== null) {
        if (match.index > lastIndex) {
            parts.push(new TextRun({ text: trimmedLine.substring(lastIndex, match.index) }));
        }
        parts.push(new TextRun({ text: match[1], bold: true }));
        lastIndex = boldRegex.lastIndex;
    }
    if (lastIndex < trimmedLine.length) {
        parts.push(new TextRun({ text: trimmedLine.substring(lastIndex) }));
    }
    
    if (parts.length === 0) {
        parts.push(new TextRun({ text: trimmedLine }));
    }
    
    children.push(new Paragraph({
        spacing: { after: 120 },
        children: parts
    }));
}

const doc = new Document({
    styles: {
        default: {
            document: {
                run: { font: "Arial", size: 24 }
            }
        },
        paragraphStyles: [
            {
                id: "Title",
                name: "Title",
                basedOn: "Normal",
                run: { size: 48, bold: true, color: "000000", font: "Arial" },
                paragraph: { spacing: { before: 0, after: 200 }, alignment: AlignmentType.CENTER }
            },
            {
                id: "Heading1",
                name: "Heading 1",
                basedOn: "Normal",
                next: "Normal",
                quickFormat: true,
                run: { size: 32, bold: true, color: "1F4E79", font: "Arial" },
                paragraph: { spacing: { before: 400, after: 200 } }
            },
            {
                id: "Heading2",
                name: "Heading 2",
                basedOn: "Normal",
                next: "Normal",
                quickFormat: true,
                run: { size: 26, bold: true, color: "2E75B6", font: "Arial" },
                paragraph: { spacing: { before: 300, after: 150 } }
            }
        ]
    },
    numbering: {
        config: [{
            reference: "bullet-list",
            levels: [{
                level: 0,
                format: LevelFormat.BULLET,
                text: "•",
                alignment: AlignmentType.LEFT,
                style: { paragraph: { indent: { left: 720, hanging: 360 } } }
            }]
        }]
    },
    sections: [{
        properties: {
            page: {
                margin: { top: 1440, right: 1440, bottom: 1440, left: 1440 }
            }
        },
        children: children
    }]
});

Packer.toBuffer(doc).then(buffer => {
    fs.writeFileSync('/tmp/review_output.docx', buffer);
    console.log('Document created successfully');
}).catch(err => {
    console.error('Error creating document:', err);
    process.exit(1);
});
'''
        
        # Write files
        with open('/tmp/review_content.txt', 'w') as f:
            f.write(review_text)
        
        with open('/tmp/contract_name.txt', 'w') as f:
            f.write(contract_name)
        
        with open('/tmp/create_review.js', 'w') as f:
            f.write(js_code)
        
        # Run node
        result = subprocess.run(
            ['node', '/tmp/create_review.js'],
            capture_output=True,
            text=True,
            timeout=60
        )
        
        if result.returncode == 0 and os.path.exists('/tmp/review_output.docx'):
            with open('/tmp/review_output.docx', 'rb') as f:
                return f.read()
        else:
            st.error(f"Error creating DOCX: {result.stderr}")
            return None
            
    except Exception as e:
        st.error(f"Error creating DOCX: {e}")
        return None

# ============================================================
# MAIN UI
# ============================================================

# Check API key
if not CLAUDE_API_KEY:
    st.error("❌ Claude API key (ANTHROPIC_API_KEY) not configured in secrets")
    st.stop()

# Input options
st.subheader("📄 Contract Input")

input_method = st.radio(
    "How would you like to provide the contract?",
    ["Upload File (PDF, DOC, DOCX)", "Paste Text", "Google Doc URL"],
    horizontal=True
)

contract_text = None
contract_name = "Uploaded Contract"

if input_method == "Upload File (PDF, DOC, DOCX)":
    uploaded_file = st.file_uploader(
        "Upload contract file",
        type=['pdf', 'doc', 'docx'],
        help="Supported formats: PDF, DOC, DOCX"
    )
    
    if uploaded_file:
        contract_name = uploaded_file.name
        
        with st.spinner(f"📄 Extracting text from {uploaded_file.name}..."):
            if uploaded_file.name.lower().endswith('.pdf'):
                contract_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.name.lower().endswith('.docx'):
                contract_text = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.lower().endswith('.doc'):
                contract_text = extract_text_from_doc(uploaded_file)
        
        if contract_text:
            st.success(f"✅ Extracted {len(contract_text):,} characters from {uploaded_file.name}")
            with st.expander("Preview extracted text"):
                st.text(contract_text[:3000] + "..." if len(contract_text) > 3000 else contract_text)

elif input_method == "Paste Text":
    contract_text = st.text_area(
        "Paste contract text",
        height=300,
        placeholder="Paste the full contract text here..."
    )
    if contract_text:
        contract_name = "Pasted Contract"
        st.info(f"📝 {len(contract_text):,} characters entered")

elif input_method == "Google Doc URL":
    google_doc_url = st.text_input(
        "Google Doc URL",
        placeholder="https://docs.google.com/document/d/..."
    )
    
    if google_doc_url:
        # Extract doc ID from URL
        import re
        match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', google_doc_url)
        if match:
            doc_id = match.group(1)
            contract_name = f"Google Doc {doc_id[:8]}..."
            
            with st.spinner("📄 Fetching Google Doc..."):
                contract_text = fetch_google_doc_content(doc_id)
            
            if contract_text:
                st.success(f"✅ Fetched {len(contract_text):,} characters")
                with st.expander("Preview extracted text"):
                    st.text(contract_text[:3000] + "..." if len(contract_text) > 3000 else contract_text)
        else:
            st.error("Invalid Google Doc URL. Please provide a valid URL.")

# Review button
st.divider()

if st.button("🔍 Review Contract", type="primary", disabled=not contract_text):
    
    # Load standards
    with st.spinner("📚 Loading Voyage contract standards..."):
        standards_text = fetch_google_doc_content(CONTRACT_STANDARDS_DOC_ID)
        
        if not standards_text:
            st.error("❌ Could not load contract standards")
            st.stop()
        
        st.success("✅ Loaded contract standards")
    
    # Call Claude
    with st.spinner("🤖 Analyzing contract with Claude AI... (this may take 1-2 minutes)"):
        review_result = call_claude_api(contract_text, standards_text)
    
    if review_result:
        st.success("✅ Review complete!")
        
        # Display review
        st.divider()
        st.header("📋 Contract Review Results")
        
        # Show the review in markdown
        st.markdown(review_result)
        
        # Create downloadable DOCX
        st.divider()
        st.subheader("📥 Download Report")
        
        with st.spinner("Creating Word document..."):
            docx_bytes = create_review_docx(review_result, contract_name)
        
        if docx_bytes:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Contract_Review_{timestamp}.docx"
            
            st.download_button(
                label="📥 Download Review as Word Document",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            # Fallback to text download
            st.download_button(
                label="📥 Download Review as Text",
                data=review_result,
                file_name=f"Contract_Review_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )
    else:
        st.error("❌ Failed to generate review")

else:
    if not contract_text:
        st.info("👆 Please provide a contract to review")

# Info section
with st.expander("ℹ️ About Contract Reviewer"):
    st.markdown("""
    ### How It Works
    
    1. **Upload or paste** your contract (PDF, DOC, DOCX, or text)
    2. The app loads **Voyage's contract standards** from Google Docs
    3. **Claude AI** analyzes the contract against the standards
    4. You receive a **detailed review** with:
       - General comments and observations
       - Section-by-section findings
       - Proposed replacement language
       - Recommendations for missing provisions
    
    ### Review Categories
    
    - Limitation of Liability
    - Work Product and Intellectual Property
    - Payment Terms
    - Indemnification
    - Confidentiality
    - Termination
    - Governing Law and Venue
    - Entity Names and Signature Blocks
    
    ### Standards Document
    
    The contract standards are maintained in a Google Doc that can be updated at any time.
    Changes to the standards take effect immediately for all future reviews.
    """)
